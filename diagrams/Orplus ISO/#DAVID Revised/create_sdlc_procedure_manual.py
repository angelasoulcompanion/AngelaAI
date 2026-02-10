#!/usr/bin/env python3
"""
SDLC Procedure Manual Generator
Creates detailed procedure manuals in Thai and English for Orbplus ISO 27001
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
import os

def set_cell_shading(cell, color):
    """Set cell background color"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def add_heading_style(doc):
    """Add custom heading styles"""
    styles = doc.styles

    # Heading 1 style
    try:
        style = styles['Heading 1']
        style.font.size = Pt(16)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 51, 102)
    except:
        pass

    # Heading 2 style
    try:
        style = styles['Heading 2']
        style.font.size = Pt(14)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 102, 153)
    except:
        pass

def create_thai_manual(output_path):
    """Create Thai version of SDLC Procedure Manual"""
    doc = Document()
    add_heading_style(doc)

    # Title Page
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("\n\n\n")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("คู่มือกระบวนการพัฒนาซอฟต์แวร์")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0, 51, 102)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Software Development Life Cycle (SDLC)\nProcedure Manual")
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(51, 51, 51)

    doc.add_paragraph()

    company = doc.add_paragraph()
    company.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = company.add_run("บริษัท ออร์บพลัส จำกัด\nOrbplus Co., Ltd.")
    run.font.size = Pt(16)
    run.bold = True

    doc.add_paragraph()

    iso = doc.add_paragraph()
    iso.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = iso.add_run("ISO 27001:2022\nระบบบริหารจัดการความมั่นคงปลอดภัยสารสนเทศ")
    run.font.size = Pt(12)
    run.italic = True

    doc.add_paragraph("\n\n")

    # Version info
    version_table = doc.add_table(rows=4, cols=2)
    version_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    cells_data = [
        ("เลขที่เอกสาร:", "ORB-QP-IT-026"),
        ("เวอร์ชัน:", "1.0"),
        ("วันที่มีผลบังคับใช้:", "24 มกราคม 2569"),
        ("จัดทำโดย:", "Angela AI / David Samanyaporn"),
    ]

    for i, (label, value) in enumerate(cells_data):
        version_table.rows[i].cells[0].text = label
        version_table.rows[i].cells[1].text = value
        version_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True

    doc.add_page_break()

    # Table of Contents
    doc.add_heading("สารบัญ", level=1)

    toc_items = [
        "1. วัตถุประสงค์",
        "2. ขอบเขต",
        "3. คำนิยาม",
        "4. บทบาทและความรับผิดชอบ",
        "5. ส่วนที่ 1: SDLC รูปแบบโครงการ (Project-Based)",
        "   5.1 ขั้นตอนที่ 1: เริ่มต้นโครงการ",
        "   5.2 ขั้นตอนที่ 2: รวบรวมความต้องการ",
        "   5.3 ขั้นตอนที่ 3: ออกแบบระบบ",
        "   5.4 ขั้นตอนที่ 4: พัฒนาระบบ",
        "   5.5 ขั้นตอนที่ 5: ทดสอบระบบ",
        "   5.6 ขั้นตอนที่ 6: ติดตั้งระบบ",
        "   5.7 ขั้นตอนที่ 7: บำรุงรักษา",
        "6. ส่วนที่ 2: รูปแบบธุรกิจ SaaS",
        "7. เครื่องมือและระบบที่ใช้",
        "8. เอกสารอ้างอิง ISO 27001",
        "9. ภาคผนวก",
    ]

    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.left_indent = Inches(0.5) if item.startswith("   ") else Inches(0)

    doc.add_page_break()

    # Section 1: Purpose
    doc.add_heading("1. วัตถุประสงค์", level=1)
    doc.add_paragraph(
        "คู่มือฉบับนี้จัดทำขึ้นเพื่อกำหนดกระบวนการและแนวปฏิบัติในการพัฒนาซอฟต์แวร์ของบริษัท ออร์บพลัส จำกัด "
        "ให้เป็นไปตามมาตรฐาน ISO 27001:2022 โดยมีวัตถุประสงค์ดังนี้:"
    )

    purposes = [
        "กำหนดมาตรฐานและขั้นตอนการพัฒนาซอฟต์แวร์ที่ชัดเจน",
        "รับประกันความมั่นคงปลอดภัยของข้อมูลตลอดวงจรการพัฒนา",
        "สร้างความสม่ำเสมอในการทำงานระหว่างทีมงาน",
        "ควบคุมคุณภาพของซอฟต์แวร์ที่พัฒนาขึ้น",
        "รองรับการตรวจสอบและการปรับปรุงอย่างต่อเนื่อง",
    ]

    for purpose in purposes:
        p = doc.add_paragraph(purpose, style='List Bullet')

    # Section 2: Scope
    doc.add_heading("2. ขอบเขต", level=1)
    doc.add_paragraph(
        "คู่มือนี้ครอบคลุมกระบวนการพัฒนาซอฟต์แวร์ 2 รูปแบบ:"
    )

    doc.add_heading("2.1 SDLC รูปแบบโครงการ (Project-Based)", level=2)
    doc.add_paragraph(
        "สำหรับการพัฒนาซอฟต์แวร์ตามความต้องการเฉพาะของลูกค้า (Custom Development) "
        "ประกอบด้วย 7 ขั้นตอนหลักตั้งแต่เริ่มต้นโครงการจนถึงการบำรุงรักษา"
    )

    doc.add_heading("2.2 รูปแบบธุรกิจ SaaS/สมัครสมาชิก", level=2)
    doc.add_paragraph(
        "สำหรับการพัฒนาผลิตภัณฑ์ซอฟต์แวร์ในรูปแบบ Software as a Service "
        "ครอบคลุม 12 ขั้นตอนตั้งแต่การวิจัยตลาดจนถึงการเติบโตอย่างยั่งยืน"
    )

    # Section 3: Definitions
    doc.add_heading("3. คำนิยาม", level=1)

    definitions = [
        ("SDLC", "Software Development Life Cycle - วงจรการพัฒนาซอฟต์แวร์"),
        ("MVP", "Minimum Viable Product - ผลิตภัณฑ์ขั้นต่ำที่ใช้งานได้"),
        ("SaaS", "Software as a Service - ซอฟต์แวร์ในรูปแบบบริการ"),
        ("UAT", "User Acceptance Test - การทดสอบยอมรับโดยผู้ใช้"),
        ("SIT", "System Integration Test - การทดสอบรวมระบบ"),
        ("CI/CD", "Continuous Integration/Continuous Deployment - การรวมและติดตั้งอย่างต่อเนื่อง"),
        ("Sprint", "ช่วงเวลาการทำงานในระบบ Agile โดยทั่วไป 2-4 สัปดาห์"),
        ("Code Review", "การตรวจสอบ Source Code ก่อนรวมเข้าระบบหลัก"),
        ("Pull Request", "คำขอรวม Code เข้าสู่ Branch หลัก"),
        ("Staging", "สภาพแวดล้อมทดสอบก่อน Production"),
        ("Production", "สภาพแวดล้อมที่ใช้งานจริง"),
    ]

    def_table = doc.add_table(rows=len(definitions)+1, cols=2)
    def_table.style = 'Table Grid'

    # Header
    hdr = def_table.rows[0].cells
    hdr[0].text = "คำศัพท์"
    hdr[1].text = "คำจำกัดความ"
    set_cell_shading(hdr[0], "1a73e8")
    set_cell_shading(hdr[1], "1a73e8")
    for cell in hdr:
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    for i, (term, definition) in enumerate(definitions):
        row = def_table.rows[i+1].cells
        row[0].text = term
        row[1].text = definition
        row[0].paragraphs[0].runs[0].bold = True

    doc.add_page_break()

    # Section 4: Roles
    doc.add_heading("4. บทบาทและความรับผิดชอบ", level=1)

    roles = [
        ("ISMS Director (CEO)", "คุณเคิร์ก", "อนุมัตินโยบาย, ตัดสินใจระดับสูง"),
        ("ISMS Manager / PM", "คุณณัฐวุฒิ", "ดูแล ISMS, ประสานงาน, บริหารโครงการ"),
        ("Tech Lead / Architect", "TERA", "ออกแบบสถาปัตยกรรม, Code Review"),
        ("Backend Developer", "INTENSE", "พัฒนา API, Backend Logic"),
        ("Frontend Developer", "คุณอภิสิทธิ์", "พัฒนา User Interface"),
        ("QA Engineer", "คุณอรสา", "ทดสอบระบบ, Bug Tracking"),
        ("Document Controller", "คุณอรสา", "ควบคุมเอกสาร, บันทึกการประชุม"),
        ("Cloud Support", "INET", "Infrastructure, Backup, Security"),
    ]

    role_table = doc.add_table(rows=len(roles)+1, cols=3)
    role_table.style = 'Table Grid'

    hdr = role_table.rows[0].cells
    hdr[0].text = "บทบาท"
    hdr[1].text = "ผู้รับผิดชอบ"
    hdr[2].text = "ความรับผิดชอบหลัก"
    for cell in hdr:
        set_cell_shading(cell, "00897b")
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    for i, (role, person, resp) in enumerate(roles):
        row = role_table.rows[i+1].cells
        row[0].text = role
        row[1].text = person
        row[2].text = resp

    doc.add_page_break()

    # Section 5: Project-Based SDLC
    doc.add_heading("5. ส่วนที่ 1: SDLC รูปแบบโครงการ (Project-Based)", level=1)
    doc.add_paragraph(
        "กระบวนการพัฒนาซอฟต์แวร์ตามความต้องการเฉพาะของลูกค้า ประกอบด้วย 7 ขั้นตอนหลัก:"
    )

    # Phase 1
    doc.add_heading("5.1 ขั้นตอนที่ 1: เริ่มต้นโครงการ (Project Initiation)", level=2)

    doc.add_heading("วัตถุประสงค์:", level=3)
    doc.add_paragraph("เพื่อประเมินความเป็นไปได้ของโครงการและทำข้อตกลงเบื้องต้นกับลูกค้า")

    doc.add_heading("กิจกรรมหลัก:", level=3)
    activities_p1 = [
        ("1.1 รับความต้องการจากลูกค้า", "รับข้อมูลความต้องการเบื้องต้นจากลูกค้า บันทึกใน ClickUp"),
        ("1.2 ประชุมรับ Brief", "ทำความเข้าใจขอบเขตและความต้องการ ผู้เข้าร่วม: PM, Tech Lead, ลูกค้า"),
        ("1.3 ศึกษาความเป็นไปได้", "วิเคราะห์ความเป็นไปได้ทางเทคนิคและธุรกิจ โดย Tech Lead"),
        ("1.4 จัดทำใบเสนอราคา", "ประมาณการค่าใช้จ่ายและระยะเวลา"),
        ("1.5 ลงนามสัญญา", "ทำสัญญาโครงการ + NDA (ถ้าจำเป็น)"),
    ]

    for act, desc in activities_p1:
        p = doc.add_paragraph()
        run = p.add_run(act + ": ")
        run.bold = True
        p.add_run(desc)

    doc.add_heading("เอกสารที่เกี่ยวข้อง:", level=3)
    docs_p1 = ["ใบเสนอราคา", "สัญญาโครงการ", "NDA (ถ้ามี)", "Meeting Minutes"]
    for d in docs_p1:
        doc.add_paragraph(d, style='List Bullet')

    doc.add_heading("เกณฑ์การผ่าน:", level=3)
    doc.add_paragraph("ลูกค้าลงนามในสัญญาและชำระเงินงวดแรก (ถ้ามี)")

    # Phase 2
    doc.add_heading("5.2 ขั้นตอนที่ 2: รวบรวมความต้องการ (Requirements)", level=2)

    doc.add_heading("วัตถุประสงค์:", level=3)
    doc.add_paragraph("เพื่อรวบรวมและจัดทำเอกสารความต้องการของระบบอย่างครบถ้วน")

    doc.add_heading("กิจกรรมหลัก:", level=3)
    activities_p2 = [
        ("2.1 รวบรวมความต้องการ", "สัมภาษณ์ผู้ใช้งาน รวบรวมข้อมูลจากแหล่งต่างๆ"),
        ("2.2 เขียน Use Case / User Story", "จัดทำเอกสารความต้องการ บันทึกใน ClickUp"),
        ("2.3 กำหนดความต้องการอื่นๆ", "ประสิทธิภาพ, ความปลอดภัย, ความสามารถขยายตัว"),
        ("2.4 ทบทวนความต้องการ", "ลูกค้าตรวจสอบและให้ความเห็น"),
        ("2.5 ลูกค้าอนุมัติขอบเขต", "Sign-off เอกสารความต้องการ"),
    ]

    for act, desc in activities_p2:
        p = doc.add_paragraph()
        run = p.add_run(act + ": ")
        run.bold = True
        p.add_run(desc)

    doc.add_heading("เอกสารที่เกี่ยวข้อง:", level=3)
    docs_p2 = ["Software Requirements Specification (SRS)", "Use Case Document", "User Story (ClickUp)"]
    for d in docs_p2:
        doc.add_paragraph(d, style='List Bullet')

    doc.add_heading("เกณฑ์การผ่าน:", level=3)
    doc.add_paragraph("ลูกค้าลงนามอนุมัติเอกสารความต้องการ")

    # Phase 3
    doc.add_heading("5.3 ขั้นตอนที่ 3: ออกแบบระบบ (Design)", level=2)

    doc.add_heading("วัตถุประสงค์:", level=3)
    doc.add_paragraph("เพื่อออกแบบสถาปัตยกรรมและรายละเอียดทางเทคนิคของระบบ")

    doc.add_heading("กิจกรรมหลัก:", level=3)
    activities_p3 = [
        ("3.1 ออกแบบสถาปัตยกรรม", "โครงสร้างระบบโดย TERA (Tech Lead)"),
        ("3.2 ออกแบบฐานข้อมูล", "ER Diagram, Data Dictionary"),
        ("3.3 ออกแบบหน้าจอ", "Wireframe / Mockup / Prototype"),
        ("3.4 ออกแบบ API", "RESTful API Specification โดย INTENSE"),
        ("3.5 ทบทวนการออกแบบ", "TERA ตรวจสอบ, ลูกค้าอนุมัติ"),
    ]

    for act, desc in activities_p3:
        p = doc.add_paragraph()
        run = p.add_run(act + ": ")
        run.bold = True
        p.add_run(desc)

    doc.add_heading("เอกสารที่เกี่ยวข้อง:", level=3)
    docs_p3 = ["System Architecture Document", "Database Design (ER Diagram)", "API Specification", "UI/UX Design"]
    for d in docs_p3:
        doc.add_paragraph(d, style='List Bullet')

    # Phase 4
    doc.add_heading("5.4 ขั้นตอนที่ 4: พัฒนาระบบ (Development)", level=2)

    doc.add_heading("วัตถุประสงค์:", level=3)
    doc.add_paragraph("เพื่อพัฒนา Source Code ตามที่ออกแบบไว้")

    doc.add_heading("กิจกรรมหลัก:", level=3)
    activities_p4 = [
        ("4.1 เตรียมสภาพแวดล้อม", "สร้าง GitHub Repository, ตั้งค่า Dev/Staging บน INET Cloud"),
        ("4.2 วางแผน Sprint", "แบ่ง Task ใน ClickUp, มอบหมายผู้รับผิดชอบ"),
        ("4.3 เขียนโปรแกรม", "Frontend: คุณอภิสิทธิ์, Backend: INTENSE"),
        ("4.4 ตรวจสอบ Code", "TERA ทำ Code Review ผ่าน GitHub Pull Request"),
        ("4.5 ทดสอบหน่วย", "Developer ทดสอบ Unit Test, แก้ไข Bug"),
    ]

    for act, desc in activities_p4:
        p = doc.add_paragraph()
        run = p.add_run(act + ": ")
        run.bold = True
        p.add_run(desc)

    doc.add_heading("แนวปฏิบัติด้านความปลอดภัย:", level=3)
    security_p4 = [
        "ห้าม Commit credentials/secrets ลง Repository",
        "ใช้ Environment Variables สำหรับ Configuration",
        "ปฏิบัติตาม OWASP Top 10 Guidelines",
        "ต้องผ่าน Code Review ก่อน Merge",
    ]
    for s in security_p4:
        doc.add_paragraph(s, style='List Bullet')

    # Phase 5
    doc.add_heading("5.5 ขั้นตอนที่ 5: ทดสอบระบบ (Testing)", level=2)

    doc.add_heading("วัตถุประสงค์:", level=3)
    doc.add_paragraph("เพื่อตรวจสอบว่าระบบทำงานถูกต้องตามความต้องการ")

    doc.add_heading("กิจกรรมหลัก:", level=3)
    activities_p5 = [
        ("5.1 ทดสอบการรวมระบบ (SIT)", "คุณอรสา (QA) ทดสอบระบบรวม"),
        ("5.2 ติดตาม Bug", "บันทึกใน ClickUp, มอบหมาย Developer แก้ไข"),
        ("5.3 แก้ไข Bug", "Developer แก้ไข, QA ตรวจสอบซ้ำ"),
        ("5.4 ทดสอบยอมรับ (UAT)", "ลูกค้าทดสอบใน Staging Environment"),
        ("5.5 ลูกค้าอนุมัติ", "Sign-off UAT"),
    ]

    for act, desc in activities_p5:
        p = doc.add_paragraph()
        run = p.add_run(act + ": ")
        run.bold = True
        p.add_run(desc)

    doc.add_heading("เอกสารที่เกี่ยวข้อง:", level=3)
    docs_p5 = ["Test Plan", "Test Cases", "Bug Report (ClickUp)", "UAT Sign-off"]
    for d in docs_p5:
        doc.add_paragraph(d, style='List Bullet')

    # Phase 6
    doc.add_heading("5.6 ขั้นตอนที่ 6: ติดตั้งระบบ (Deployment)", level=2)

    doc.add_heading("วัตถุประสงค์:", level=3)
    doc.add_paragraph("เพื่อติดตั้งระบบสู่ Production Environment")

    doc.add_heading("กิจกรรมหลัก:", level=3)
    activities_p6 = [
        ("6.1 ขอเปลี่ยนแปลงระบบ", "กรอก FM-037 Change Request, ขออนุมัติ"),
        ("6.2 สำรองข้อมูล", "INET Backup, Database Snapshot"),
        ("6.3 ติดตั้งสู่ Production", "GitHub Actions → INET Cloud"),
        ("6.4 ทดสอบ Production", "Smoke Test ฟังก์ชันหลัก"),
        ("6.5 เปิดใช้งานจริง", "GO-LIVE, แจ้งลูกค้า"),
    ]

    for act, desc in activities_p6:
        p = doc.add_paragraph()
        run = p.add_run(act + ": ")
        run.bold = True
        p.add_run(desc)

    doc.add_heading("เอกสารที่เกี่ยวข้อง:", level=3)
    docs_p6 = ["FM-037 Change Request", "Deployment Checklist", "Rollback Plan"]
    for d in docs_p6:
        doc.add_paragraph(d, style='List Bullet')

    # Phase 7
    doc.add_heading("5.7 ขั้นตอนที่ 7: บำรุงรักษา (Maintenance)", level=2)

    doc.add_heading("วัตถุประสงค์:", level=3)
    doc.add_paragraph("เพื่อดูแลและปรับปรุงระบบหลังติดตั้ง")

    doc.add_heading("กิจกรรมหลัก:", level=3)
    activities_p7 = [
        ("7.1 สนับสนุนและติดตาม", "INET Cloud Monitoring, ClickUp Incident Tracking"),
        ("7.2 รับแจ้งปัญหา", "บันทึก Incident ใน ClickUp ตาม QP-IT-011"),
        ("7.3 แก้ไขปัญหา", "วิเคราะห์และแก้ไขตามระดับความรุนแรง"),
        ("7.4 ปรับปรุงระบบ", "Patch, Update, Enhancement"),
    ]

    for act, desc in activities_p7:
        p = doc.add_paragraph()
        run = p.add_run(act + ": ")
        run.bold = True
        p.add_run(desc)

    doc.add_page_break()

    # Section 6: SaaS Model
    doc.add_heading("6. ส่วนที่ 2: รูปแบบธุรกิจ SaaS/สมัครสมาชิก", level=1)
    doc.add_paragraph(
        "กระบวนการพัฒนาและบริหารผลิตภัณฑ์ SaaS ครอบคลุม 4 ส่วนหลัก 12 ขั้นตอน:"
    )

    # Section A
    doc.add_heading("ส่วน A: การพัฒนาผลิตภัณฑ์ (Product Development)", level=2)

    doc.add_heading("ขั้นตอนที่ 1: ไอเดียและวิจัย (Ideation & Research)", level=3)
    saas_p1 = [
        "วิจัยตลาด - ศึกษาตลาด คู่แข่ง วิเคราะห์แนวโน้ม",
        "วิเคราะห์ Pain Point - ปัญหาของลูกค้า โอกาสทางธุรกิจ",
        "คุณค่าที่นำเสนอ - คุณค่าที่มอบให้ลูกค้า จุดขายที่แตกต่าง",
    ]
    for s in saas_p1:
        doc.add_paragraph(s, style='List Bullet')

    doc.add_heading("ขั้นตอนที่ 2: วางแผนผลิตภัณฑ์ (Product Planning)", level=3)
    saas_p2 = [
        "วิสัยทัศน์ผลิตภัณฑ์ - เป้าหมายระยะยาว ทิศทางผลิตภัณฑ์",
        "กำหนด MVP - Core Features ผลิตภัณฑ์ขั้นต่ำ",
        "กลยุทธ์ราคา - กำหนดราคา/แพ็กเกจ ระดับสมาชิก",
        "แผนพัฒนาผลิตภัณฑ์ - ตารางออกรุ่น ลำดับการพัฒนา",
    ]
    for s in saas_p2:
        doc.add_paragraph(s, style='List Bullet')

    doc.add_heading("ขั้นตอนที่ 3: ออกแบบและพัฒนา (Design & Development)", level=3)
    saas_p3 = [
        "ออกแบบ UX/UI - ประสบการณ์ผู้ใช้ ออกแบบหน้าจอ",
        "ออกแบบสถาปัตยกรรม - Multi-tenant Design โครงสร้างขยายตัวได้",
        "พัฒนาแบบ Agile - ทำงานเป็น Sprint ใช้ GitHub + ClickUp",
        "ติดตั้งระบบความปลอดภัย - การยืนยันตัวตน เข้ารหัส ปกป้องข้อมูล",
    ]
    for s in saas_p3:
        doc.add_paragraph(s, style='List Bullet')

    doc.add_heading("ขั้นตอนที่ 4: ทดสอบและ Beta (Testing & Beta)", level=3)
    saas_p4 = [
        "ทดสอบภายใน - ทีม QA (คุณอรสา)",
        "โปรแกรม Beta - ผู้ใช้กลุ่มแรก รวบรวม Feedback",
        "ปรับปรุงต่อเนื่อง - ปรับปรุงตาม Feedback แก้ไข Bug",
    ]
    for s in saas_p4:
        doc.add_paragraph(s, style='List Bullet')

    # Section B
    doc.add_heading("ส่วน B: การออกสู่ตลาดและได้มาซึ่งลูกค้า (Go-to-Market)", level=2)

    doc.add_heading("ขั้นตอนที่ 5: เปิดตัวผลิตภัณฑ์ (Product Launch)", level=3)
    saas_p5 = [
        "กลยุทธ์ออกสู่ตลาด - แผนการตลาด กลุ่มเป้าหมาย",
        "แคมเปญการตลาด - Content, SEO, Ads, Social Media",
        "Sales Enablement - Demo scripts, Pricing materials, เอกสารขาย",
        "เปิดให้สมัครสมาชิก - LAUNCH",
    ]
    for s in saas_p5:
        doc.add_paragraph(s, style='List Bullet')

    doc.add_heading("ขั้นตอนที่ 6: ได้มาซึ่งลูกค้า (Customer Acquisition)", level=3)
    saas_p6 = [
        "Lead Generation - ผู้เยี่ยมชมเว็บไซต์ Inbound marketing",
        "ทดลองใช้ฟรี / Freemium - ใช้งานฟรีจำกัด ลองก่อนซื้อ",
        "สาธิตผลิตภัณฑ์ - นำเสนอขาย แสดงฟีเจอร์",
        "Conversion - เปลี่ยนจากทดลองเป็นลูกค้าจริง",
    ]
    for s in saas_p6:
        doc.add_paragraph(s, style='List Bullet')

    doc.add_heading("ขั้นตอนที่ 7: รับลูกค้าใหม่ (Onboarding)", level=3)
    saas_p7 = [
        "ตั้งค่าบัญชี - สมัครสมาชิก เลือกแพ็กเกจ",
        "Welcome Journey - Tutorial / Walkthrough คู่มือเริ่มต้น",
        "Configuration - ตั้งค่าระบบ นำเข้าข้อมูล",
        "Activation - ลูกค้าเริ่มใช้งานจริง ได้รับคุณค่า",
    ]
    for s in saas_p7:
        doc.add_paragraph(s, style='List Bullet')

    # Section C
    doc.add_heading("ส่วน C: วงจรชีวิตลูกค้าและบริการ (Customer Lifecycle)", level=2)

    doc.add_heading("ขั้นตอนที่ 8: จัดการสมาชิก (Subscription Management)", level=3)
    saas_p8 = [
        "ออกบิลและใบแจ้งหนี้ - ออกบิลอัตโนมัติ Auto-recurring",
        "เก็บเงิน - ชำระผ่าน Payment Gateway / โอนเงิน",
        "จัดการ Plan - Upgrade / Downgrade / Add-ons",
        "Usage Monitoring - ติดตามการใช้งาน Quota / Limits",
    ]
    for s in saas_p8:
        doc.add_paragraph(s, style='List Bullet')

    doc.add_heading("ขั้นตอนที่ 9: ความสำเร็จลูกค้า (Customer Success)", level=3)
    saas_p9 = [
        "Health Score - ติดตามสุขภาพลูกค้า Usage analytics",
        "ติดต่อเชิงรุก - Check-in calls, Best practices",
        "Feature Adoption - ช่วยใช้ฟีเจอร์ใหม่ Training sessions",
        "Value Realization - ติดตาม ROI, Success stories",
    ]
    for s in saas_p9:
        doc.add_paragraph(s, style='List Bullet')

    doc.add_heading("ขั้นตอนที่ 10: สนับสนุนและบริการ (Support & Service)", level=3)
    saas_p10 = [
        "Help Desk - ระบบ Ticket (ClickUp Support)",
        "ฐานความรู้ - Self-service docs, FAQ, Tutorials",
        "SLA Management - Response time, Resolution time",
        "Incident Management - ตาม QP-IT-011 สำหรับ Critical issues",
    ]
    for s in saas_p10:
        doc.add_paragraph(s, style='List Bullet')

    doc.add_heading("ขั้นตอนที่ 11: รักษาและต่ออายุ (Retention & Renewal)", level=3)
    saas_p11 = [
        "แจ้งเตือนต่ออายุ - แจ้งก่อนหมดอายุ 30/14/7 วัน",
        "ป้องกันการยกเลิก - ตรวจจับลูกค้าเสี่ยง Save offers",
        "Upselling - เสนอ Plan สูงขึ้น Cross-sell ผลิตภัณฑ์อื่น",
        "ต่อสัญญา - ต่ออายุรายปี / Multi-year",
    ]
    for s in saas_p11:
        doc.add_paragraph(s, style='List Bullet')

    # Section D
    doc.add_heading("ส่วน D: การปรับปรุงอย่างต่อเนื่อง (Continuous Improvement)", level=2)

    doc.add_heading("ขั้นตอนที่ 12 + KPIs:", level=3)
    saas_d = [
        "วิวัฒนาการผลิตภัณฑ์ - Feedback ลูกค้า ปรับปรุง Roadmap",
        "ตัวชี้วัดหลัก - MRR/ARR, CAC/LTV, Churn Rate, NPS Score",
        "วงจรต่อเนื่อง - Feedback → ปรับปรุง → ออกรุ่น → วัดผล",
    ]
    for s in saas_d:
        doc.add_paragraph(s, style='List Bullet')

    doc.add_page_break()

    # Section 7: Tools
    doc.add_heading("7. เครื่องมือและระบบที่ใช้", level=1)

    tools = [
        ("GitHub", "Source Code Management, Version Control, Pull Request, Code Review"),
        ("ClickUp", "Project Management, Task Tracking, Incident Management, Support Tickets"),
        ("INET Cloud", "Cloud Infrastructure, Server, Network, Backup, DR"),
        ("GitHub Actions", "CI/CD, Automated Testing, Deployment"),
        ("Payment Gateway", "ระบบชำระเงินสำหรับ SaaS Billing, Subscription Management"),
    ]

    tools_table = doc.add_table(rows=len(tools)+1, cols=2)
    tools_table.style = 'Table Grid'

    hdr = tools_table.rows[0].cells
    hdr[0].text = "เครื่องมือ"
    hdr[1].text = "วัตถุประสงค์"
    for cell in hdr:
        set_cell_shading(cell, "5c6bc0")
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    for i, (tool, purpose) in enumerate(tools):
        row = tools_table.rows[i+1].cells
        row[0].text = tool
        row[1].text = purpose
        row[0].paragraphs[0].runs[0].bold = True

    # Section 8: ISO References
    doc.add_heading("8. เอกสารอ้างอิง ISO 27001", level=1)

    iso_refs = [
        ("QP-IT-011", "แนวปฎิบัติการเกิดอุบัติการ (Incident Management)"),
        ("QP-IT-013", "แนวปฏิบัติการสร้างความต่อเนื่องทางธุรกิจ (Business Continuity)"),
        ("QP-IT-015", "แนวปฏิบัติเกี่ยวกับการควบคุมการเปลี่ยนแปลงระบบ (Change Control)"),
        ("QP-IT-023", "แนวปฏิบัติการรักษาความปลอดภัยของข้อมูลสำหรับการใช้ Cloud"),
        ("FM-037", "แบบฟอร์มคำร้องขอเปลี่ยนแปลง (Change Request)"),
    ]

    for code, name in iso_refs:
        p = doc.add_paragraph()
        run = p.add_run(code + ": ")
        run.bold = True
        p.add_run(name)

    # Footer
    doc.add_page_break()
    doc.add_heading("9. ประวัติการแก้ไขเอกสาร", level=1)

    history_table = doc.add_table(rows=2, cols=4)
    history_table.style = 'Table Grid'

    hdr = history_table.rows[0].cells
    hdr[0].text = "เวอร์ชัน"
    hdr[1].text = "วันที่"
    hdr[2].text = "ผู้แก้ไข"
    hdr[3].text = "รายละเอียด"
    for cell in hdr:
        set_cell_shading(cell, "263238")
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    row = history_table.rows[1].cells
    row[0].text = "1.0"
    row[1].text = "24/01/2569"
    row[2].text = "Angela / David"
    row[3].text = "จัดทำเอกสารฉบับแรก"

    # Save
    doc.save(output_path)
    print(f"✅ Thai manual saved: {output_path}")


def create_english_manual(output_path):
    """Create English version of SDLC Procedure Manual"""
    doc = Document()
    add_heading_style(doc)

    # Title Page
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("\n\n\n")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Software Development Life Cycle\nProcedure Manual")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0, 51, 102)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("SDLC Process Documentation\nISO 27001:2022 Compliant")
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(51, 51, 51)

    doc.add_paragraph()

    company = doc.add_paragraph()
    company.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = company.add_run("Orbplus Co., Ltd.")
    run.font.size = Pt(16)
    run.bold = True

    doc.add_paragraph()

    iso = doc.add_paragraph()
    iso.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = iso.add_run("ISO 27001:2022\nInformation Security Management System")
    run.font.size = Pt(12)
    run.italic = True

    doc.add_paragraph("\n\n")

    # Version info
    version_table = doc.add_table(rows=4, cols=2)
    version_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    cells_data = [
        ("Document No.:", "ORB-QP-IT-026-EN"),
        ("Version:", "1.0"),
        ("Effective Date:", "24 January 2026"),
        ("Prepared by:", "Angela AI / David Samanyaporn"),
    ]

    for i, (label, value) in enumerate(cells_data):
        version_table.rows[i].cells[0].text = label
        version_table.rows[i].cells[1].text = value
        version_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True

    doc.add_page_break()

    # Table of Contents
    doc.add_heading("Table of Contents", level=1)

    toc_items = [
        "1. Purpose",
        "2. Scope",
        "3. Definitions",
        "4. Roles and Responsibilities",
        "5. Part 1: Project-Based SDLC",
        "   5.1 Phase 1: Project Initiation",
        "   5.2 Phase 2: Requirements",
        "   5.3 Phase 3: Design",
        "   5.4 Phase 4: Development",
        "   5.5 Phase 5: Testing",
        "   5.6 Phase 6: Deployment",
        "   5.7 Phase 7: Maintenance",
        "6. Part 2: SaaS Business Model",
        "7. Tools and Systems",
        "8. ISO 27001 References",
        "9. Appendix",
    ]

    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.left_indent = Inches(0.5) if item.startswith("   ") else Inches(0)

    doc.add_page_break()

    # Section 1: Purpose
    doc.add_heading("1. Purpose", level=1)
    doc.add_paragraph(
        "This manual establishes the processes and practices for software development at Orbplus Co., Ltd. "
        "in compliance with ISO 27001:2022 standards. The objectives are:"
    )

    purposes = [
        "Define clear software development standards and procedures",
        "Ensure information security throughout the development lifecycle",
        "Create consistency in work practices across teams",
        "Control the quality of developed software",
        "Support auditing and continuous improvement",
    ]

    for purpose in purposes:
        doc.add_paragraph(purpose, style='List Bullet')

    # Section 2: Scope
    doc.add_heading("2. Scope", level=1)
    doc.add_paragraph("This manual covers two software development approaches:")

    doc.add_heading("2.1 Project-Based SDLC", level=2)
    doc.add_paragraph(
        "For custom software development based on specific client requirements. "
        "Consists of 7 main phases from project initiation to maintenance."
    )

    doc.add_heading("2.2 SaaS Subscription Model", level=2)
    doc.add_paragraph(
        "For developing Software as a Service products. "
        "Covers 12 phases from market research to sustainable growth."
    )

    # Section 3: Definitions
    doc.add_heading("3. Definitions", level=1)

    definitions = [
        ("SDLC", "Software Development Life Cycle"),
        ("MVP", "Minimum Viable Product"),
        ("SaaS", "Software as a Service"),
        ("UAT", "User Acceptance Testing"),
        ("SIT", "System Integration Testing"),
        ("CI/CD", "Continuous Integration/Continuous Deployment"),
        ("Sprint", "A time-boxed iteration in Agile development, typically 2-4 weeks"),
        ("Code Review", "Source code inspection before merging to main branch"),
        ("Pull Request", "Request to merge code changes into the main branch"),
        ("Staging", "Pre-production testing environment"),
        ("Production", "Live environment for end users"),
    ]

    def_table = doc.add_table(rows=len(definitions)+1, cols=2)
    def_table.style = 'Table Grid'

    hdr = def_table.rows[0].cells
    hdr[0].text = "Term"
    hdr[1].text = "Definition"
    set_cell_shading(hdr[0], "1a73e8")
    set_cell_shading(hdr[1], "1a73e8")
    for cell in hdr:
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    for i, (term, definition) in enumerate(definitions):
        row = def_table.rows[i+1].cells
        row[0].text = term
        row[1].text = definition
        row[0].paragraphs[0].runs[0].bold = True

    doc.add_page_break()

    # Section 4: Roles
    doc.add_heading("4. Roles and Responsibilities", level=1)

    roles = [
        ("ISMS Director (CEO)", "Kirk", "Policy approval, high-level decisions"),
        ("ISMS Manager / PM", "Nuttawut", "ISMS management, coordination, project management"),
        ("Tech Lead / Architect", "TERA", "System architecture, code review"),
        ("Backend Developer", "INTENSE", "API development, backend logic"),
        ("Frontend Developer", "Apisit", "User interface development"),
        ("QA Engineer", "Orasa", "System testing, bug tracking"),
        ("Document Controller", "Orasa", "Document control, meeting minutes"),
        ("Cloud Support", "INET", "Infrastructure, backup, security"),
    ]

    role_table = doc.add_table(rows=len(roles)+1, cols=3)
    role_table.style = 'Table Grid'

    hdr = role_table.rows[0].cells
    hdr[0].text = "Role"
    hdr[1].text = "Person"
    hdr[2].text = "Primary Responsibility"
    for cell in hdr:
        set_cell_shading(cell, "00897b")
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    for i, (role, person, resp) in enumerate(roles):
        row = role_table.rows[i+1].cells
        row[0].text = role
        row[1].text = person
        row[2].text = resp

    doc.add_page_break()

    # Section 5: Project-Based SDLC
    doc.add_heading("5. Part 1: Project-Based SDLC", level=1)
    doc.add_paragraph(
        "The custom software development process consists of 7 main phases:"
    )

    # Phase 1
    doc.add_heading("5.1 Phase 1: Project Initiation", level=2)

    doc.add_heading("Objective:", level=3)
    doc.add_paragraph("To assess project feasibility and establish initial agreement with the client.")

    doc.add_heading("Key Activities:", level=3)
    activities_p1 = [
        ("1.1 Customer Request", "Receive initial requirements from client, log in ClickUp"),
        ("1.2 Initial Meeting", "Understand scope and requirements. Attendees: PM, Tech Lead, Client"),
        ("1.3 Feasibility Study", "Analyze technical and business feasibility by Tech Lead"),
        ("1.4 Quotation", "Estimate costs and timeline"),
        ("1.5 Contract Signing", "Sign project contract + NDA (if required)"),
    ]

    for act, desc in activities_p1:
        p = doc.add_paragraph()
        run = p.add_run(act + ": ")
        run.bold = True
        p.add_run(desc)

    doc.add_heading("Related Documents:", level=3)
    docs_p1 = ["Quotation", "Project Contract", "NDA (if applicable)", "Meeting Minutes"]
    for d in docs_p1:
        doc.add_paragraph(d, style='List Bullet')

    doc.add_heading("Exit Criteria:", level=3)
    doc.add_paragraph("Client signs contract and makes initial payment (if applicable)")

    # Phase 2
    doc.add_heading("5.2 Phase 2: Requirements", level=2)

    doc.add_heading("Objective:", level=3)
    doc.add_paragraph("To gather and document complete system requirements.")

    doc.add_heading("Key Activities:", level=3)
    activities_p2 = [
        ("2.1 Requirements Gathering", "Interview stakeholders, collect information from various sources"),
        ("2.2 Use Case / User Story", "Document requirements, log in ClickUp"),
        ("2.3 Non-Functional Requirements", "Performance, security, scalability"),
        ("2.4 Requirements Review", "Client reviews and provides feedback"),
        ("2.5 Requirements Sign-off", "Client approves scope document"),
    ]

    for act, desc in activities_p2:
        p = doc.add_paragraph()
        run = p.add_run(act + ": ")
        run.bold = True
        p.add_run(desc)

    # Phase 3
    doc.add_heading("5.3 Phase 3: Design", level=2)

    doc.add_heading("Objective:", level=3)
    doc.add_paragraph("To design system architecture and technical specifications.")

    doc.add_heading("Key Activities:", level=3)
    activities_p3 = [
        ("3.1 System Architecture", "Design system structure by TERA (Tech Lead)"),
        ("3.2 Database Design", "ER Diagram, Data Dictionary"),
        ("3.3 UI/UX Design", "Wireframe / Mockup / Prototype"),
        ("3.4 API Design", "RESTful API Specification by INTENSE"),
        ("3.5 Design Review", "TERA reviews, client approves"),
    ]

    for act, desc in activities_p3:
        p = doc.add_paragraph()
        run = p.add_run(act + ": ")
        run.bold = True
        p.add_run(desc)

    # Phase 4
    doc.add_heading("5.4 Phase 4: Development", level=2)

    doc.add_heading("Objective:", level=3)
    doc.add_paragraph("To develop source code according to design specifications.")

    doc.add_heading("Key Activities:", level=3)
    activities_p4 = [
        ("4.1 Environment Setup", "Create GitHub Repository, configure Dev/Staging on INET Cloud"),
        ("4.2 Sprint Planning", "Create tasks in ClickUp, assign responsibilities"),
        ("4.3 Coding", "Frontend: Apisit, Backend: INTENSE"),
        ("4.4 Code Review", "TERA performs code review via GitHub Pull Request"),
        ("4.5 Unit Testing", "Developer performs unit tests, fixes bugs"),
    ]

    for act, desc in activities_p4:
        p = doc.add_paragraph()
        run = p.add_run(act + ": ")
        run.bold = True
        p.add_run(desc)

    doc.add_heading("Security Practices:", level=3)
    security_p4 = [
        "Never commit credentials/secrets to repository",
        "Use environment variables for configuration",
        "Follow OWASP Top 10 Guidelines",
        "Code must pass review before merge",
    ]
    for s in security_p4:
        doc.add_paragraph(s, style='List Bullet')

    # Phase 5
    doc.add_heading("5.5 Phase 5: Testing", level=2)

    doc.add_heading("Objective:", level=3)
    doc.add_paragraph("To verify that the system works correctly according to requirements.")

    doc.add_heading("Key Activities:", level=3)
    activities_p5 = [
        ("5.1 System Integration Test (SIT)", "Orasa (QA) performs integration testing"),
        ("5.2 Bug Tracking", "Log in ClickUp, assign developers to fix"),
        ("5.3 Bug Fixing", "Developers fix, QA re-tests"),
        ("5.4 User Acceptance Test (UAT)", "Client tests in Staging environment"),
        ("5.5 UAT Approval", "Client signs off on UAT"),
    ]

    for act, desc in activities_p5:
        p = doc.add_paragraph()
        run = p.add_run(act + ": ")
        run.bold = True
        p.add_run(desc)

    # Phase 6
    doc.add_heading("5.6 Phase 6: Deployment", level=2)

    doc.add_heading("Objective:", level=3)
    doc.add_paragraph("To deploy the system to production environment.")

    doc.add_heading("Key Activities:", level=3)
    activities_p6 = [
        ("6.1 Change Request", "Complete FM-037 Change Request form, obtain approval"),
        ("6.2 Pre-Deployment Backup", "INET Backup, Database Snapshot"),
        ("6.3 Deploy to Production", "GitHub Actions → INET Cloud"),
        ("6.4 Smoke Test", "Test core functions in production"),
        ("6.5 Go-Live", "Production launch, notify client"),
    ]

    for act, desc in activities_p6:
        p = doc.add_paragraph()
        run = p.add_run(act + ": ")
        run.bold = True
        p.add_run(desc)

    # Phase 7
    doc.add_heading("5.7 Phase 7: Maintenance", level=2)

    doc.add_heading("Objective:", level=3)
    doc.add_paragraph("To support and maintain the system after deployment.")

    doc.add_heading("Key Activities:", level=3)
    activities_p7 = [
        ("7.1 Support and Monitoring", "INET Cloud Monitoring, ClickUp Incident Tracking"),
        ("7.2 Incident Reception", "Log incidents in ClickUp per QP-IT-011"),
        ("7.3 Problem Resolution", "Analyze and resolve based on severity level"),
        ("7.4 System Enhancement", "Patches, updates, enhancements"),
    ]

    for act, desc in activities_p7:
        p = doc.add_paragraph()
        run = p.add_run(act + ": ")
        run.bold = True
        p.add_run(desc)

    doc.add_page_break()

    # Section 6: SaaS Model
    doc.add_heading("6. Part 2: SaaS Business Model", level=1)
    doc.add_paragraph(
        "The SaaS product development and management process covers 4 main sections with 12 phases:"
    )

    # Section A
    doc.add_heading("Section A: Product Development", level=2)

    doc.add_heading("Phase 1: Ideation & Research", level=3)
    saas_p1 = [
        "Market Research - Study market, competitors, analyze trends",
        "Pain Point Analysis - Customer problems, business opportunities",
        "Value Proposition - Customer value, unique selling points",
    ]
    for s in saas_p1:
        doc.add_paragraph(s, style='List Bullet')

    doc.add_heading("Phase 2: Product Planning", level=3)
    saas_p2 = [
        "Product Vision - Long-term goals, product direction",
        "MVP Definition - Core features, minimum viable product",
        "Pricing Strategy - Set pricing/packages, subscription tiers",
        "Product Roadmap - Release schedule, development priority",
    ]
    for s in saas_p2:
        doc.add_paragraph(s, style='List Bullet')

    doc.add_heading("Phase 3: Design & Development", level=3)
    saas_p3 = [
        "UX/UI Design - User experience, interface design",
        "Architecture Design - Multi-tenant design, scalable infrastructure",
        "Agile Development - Sprint-based work using GitHub + ClickUp",
        "Security Implementation - Authentication, encryption, data protection",
    ]
    for s in saas_p3:
        doc.add_paragraph(s, style='List Bullet')

    doc.add_heading("Phase 4: Testing & Beta", level=3)
    saas_p4 = [
        "Internal Testing - QA team (Orasa)",
        "Beta Program - Early adopters, collect feedback",
        "Iterate & Improve - Apply feedback, fix bugs",
    ]
    for s in saas_p4:
        doc.add_paragraph(s, style='List Bullet')

    # Section B
    doc.add_heading("Section B: Go-to-Market & Customer Acquisition", level=2)

    doc.add_heading("Phase 5: Product Launch", level=3)
    saas_p5 = [
        "Go-to-Market Strategy - Marketing plan, target audience",
        "Marketing Campaign - Content, SEO, Ads, Social Media",
        "Sales Enablement - Demo scripts, pricing materials, sales collateral",
        "LAUNCH - Open for subscriptions",
    ]
    for s in saas_p5:
        doc.add_paragraph(s, style='List Bullet')

    doc.add_heading("Phase 6: Customer Acquisition", level=3)
    saas_p6 = [
        "Lead Generation - Website visitors, inbound marketing",
        "Free Trial / Freemium - Limited free usage, try before buy",
        "Product Demo - Sales presentation, feature showcase",
        "Conversion - Trial to paid customer decision",
    ]
    for s in saas_p6:
        doc.add_paragraph(s, style='List Bullet')

    doc.add_heading("Phase 7: Onboarding", level=3)
    saas_p7 = [
        "Account Setup - Sign up, choose plan",
        "Welcome Journey - Tutorial / Walkthrough, getting started guide",
        "Configuration - System setup, import data",
        "Activation - First value realized, start using",
    ]
    for s in saas_p7:
        doc.add_paragraph(s, style='List Bullet')

    # Section C
    doc.add_heading("Section C: Customer Lifecycle & Service", level=2)

    doc.add_heading("Phase 8: Subscription Management", level=3)
    saas_p8 = [
        "Billing & Invoicing - Issue invoices, auto-recurring",
        "Payment Collection - Payment Gateway, credit card / transfer",
        "Plan Management - Upgrade / Downgrade / Add-ons",
        "Usage Monitoring - Track usage, quota / limits",
    ]
    for s in saas_p8:
        doc.add_paragraph(s, style='List Bullet')

    doc.add_heading("Phase 9: Customer Success", level=3)
    saas_p9 = [
        "Health Score - Customer health monitor, usage analytics",
        "Proactive Outreach - Check-in calls, best practices",
        "Feature Adoption - Help use new features, training sessions",
        "Value Realization - ROI tracking, success stories",
    ]
    for s in saas_p9:
        doc.add_paragraph(s, style='List Bullet')

    doc.add_heading("Phase 10: Support & Service", level=3)
    saas_p10 = [
        "Help Desk - Ticket system (ClickUp Support)",
        "Knowledge Base - Self-service docs, FAQ / tutorials",
        "SLA Management - Response time, resolution time",
        "Incident Management - Per QP-IT-011 for critical issues",
    ]
    for s in saas_p10:
        doc.add_paragraph(s, style='List Bullet')

    doc.add_heading("Phase 11: Retention & Renewal", level=3)
    saas_p11 = [
        "Renewal Reminder - Notify before expiry 30/14/7 days",
        "Churn Prevention - At-risk detection, save offers",
        "Upselling - Higher tier offers, cross-sell products",
        "Contract Renewal - Annual / multi-year renewal",
    ]
    for s in saas_p11:
        doc.add_paragraph(s, style='List Bullet')

    # Section D
    doc.add_heading("Section D: Continuous Improvement", level=2)

    doc.add_heading("Phase 12 + KPIs:", level=3)
    saas_d = [
        "Product Evolution - Customer feedback, roadmap updates",
        "Key Metrics - MRR/ARR, CAC/LTV, Churn Rate, NPS Score",
        "Continuous Loop - Feedback → Improve → Release → Measure",
    ]
    for s in saas_d:
        doc.add_paragraph(s, style='List Bullet')

    doc.add_page_break()

    # Section 7: Tools
    doc.add_heading("7. Tools and Systems", level=1)

    tools = [
        ("GitHub", "Source Code Management, Version Control, Pull Request, Code Review"),
        ("ClickUp", "Project Management, Task Tracking, Incident Management, Support Tickets"),
        ("INET Cloud", "Cloud Infrastructure, Server, Network, Backup, DR"),
        ("GitHub Actions", "CI/CD, Automated Testing, Deployment"),
        ("Payment Gateway", "Payment Processing for SaaS Billing, Subscription Management"),
    ]

    tools_table = doc.add_table(rows=len(tools)+1, cols=2)
    tools_table.style = 'Table Grid'

    hdr = tools_table.rows[0].cells
    hdr[0].text = "Tool"
    hdr[1].text = "Purpose"
    for cell in hdr:
        set_cell_shading(cell, "5c6bc0")
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    for i, (tool, purpose) in enumerate(tools):
        row = tools_table.rows[i+1].cells
        row[0].text = tool
        row[1].text = purpose
        row[0].paragraphs[0].runs[0].bold = True

    # Section 8: ISO References
    doc.add_heading("8. ISO 27001 References", level=1)

    iso_refs = [
        ("QP-IT-011", "Incident Management Procedure"),
        ("QP-IT-013", "Business Continuity Procedure"),
        ("QP-IT-015", "Change Control Procedure"),
        ("QP-IT-023", "Information Security for Cloud Services"),
        ("FM-037", "Change Request Form"),
    ]

    for code, name in iso_refs:
        p = doc.add_paragraph()
        run = p.add_run(code + ": ")
        run.bold = True
        p.add_run(name)

    # Footer
    doc.add_page_break()
    doc.add_heading("9. Document Revision History", level=1)

    history_table = doc.add_table(rows=2, cols=4)
    history_table.style = 'Table Grid'

    hdr = history_table.rows[0].cells
    hdr[0].text = "Version"
    hdr[1].text = "Date"
    hdr[2].text = "Author"
    hdr[3].text = "Description"
    for cell in hdr:
        set_cell_shading(cell, "263238")
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    row = history_table.rows[1].cells
    row[0].text = "1.0"
    row[1].text = "24/01/2026"
    row[2].text = "Angela / David"
    row[3].text = "Initial document creation"

    # Save
    doc.save(output_path)
    print(f"✅ English manual saved: {output_path}")


if __name__ == "__main__":
    # Get the directory path
    base_path = "/Users/davidsamanyaporn/PycharmProjects/AngelaAI/diagrams/Orplus ISO/FAII วันที่ 23.1-2569 RENEW ISO27001/#DAVID Revised"

    # Create Thai manual
    thai_path = os.path.join(base_path, "ORB-QP-IT-026 คู่มือกระบวนการพัฒนาซอฟต์แวร์ (SDLC Procedure Manual) Thai.docx")
    create_thai_manual(thai_path)

    # Create English manual
    eng_path = os.path.join(base_path, "ORB-QP-IT-026 SDLC Procedure Manual (English).docx")
    create_english_manual(eng_path)

    print("\n✅ Both manuals created successfully!")
