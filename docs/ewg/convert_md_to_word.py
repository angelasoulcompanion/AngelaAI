#!/usr/bin/env python3
"""
Convert Markdown to Word with PNG images
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE_DIR = Path(__file__).parent
MD_FILE = BASE_DIR / "EWG_Phase0_Data_Governance_Framework.md"
PNG_DIR = BASE_DIR / "diagrams" / "png"
OUTPUT_FILE = BASE_DIR / "EWG_Phase0_Data_Governance_Framework_Full.docx"

# Map SVG references to PNG files
IMAGE_MAP = {
    "01_data_governance_4_pillars.drawio.svg": "01_data_governance_4_pillars.png",
    "02_data_policy_hierarchy.drawio.svg": "02_data_policy_hierarchy.png",
    "03_time_of_data_capture_flow.drawio.svg": "03_time_of_data_capture_flow.png",
    "04_dama_wheel.drawio.svg": "04_dama_wheel.png",
    "05_period_end_close_timeline.drawio.svg": "05_period_end_close_timeline.png",
    "06_data_freeze_policy.drawio.svg": "06_data_freeze_policy.png",
    "07_ssot_architecture.drawio.svg": "07_ssot_architecture.png",
    "08_data_consolidation_roadmap.drawio.svg": "08_data_consolidation_roadmap.png",
    "09_data_close_cycle_policy.drawio.svg": "09_data_close_cycle_policy.png",
}


def add_hyperlink(paragraph, url, text):
    """Add hyperlink to paragraph"""
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Set color to blue
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0563C1')
    rPr.append(color)

    # Underline
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)
    return hyperlink


def parse_table(lines, start_idx):
    """Parse markdown table and return rows"""
    rows = []
    idx = start_idx

    while idx < len(lines):
        line = lines[idx].strip()
        if not line.startswith('|'):
            break

        # Skip separator line
        if re.match(r'^\|[-:\s|]+\|$', line):
            idx += 1
            continue

        # Parse cells
        cells = [cell.strip() for cell in line.split('|')[1:-1]]
        rows.append(cells)
        idx += 1

    return rows, idx


def add_table_to_doc(doc, rows):
    """Add table to document"""
    if not rows:
        return

    num_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = 'Table Grid'

    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_data in enumerate(row_data):
            if col_idx < len(row_data):
                cell = table.rows[row_idx].cells[col_idx]
                # Clean markdown formatting
                text = re.sub(r'\*\*([^*]+)\*\*', r'\1', cell_data)  # Bold
                text = re.sub(r'\*([^*]+)\*', r'\1', text)  # Italic
                text = re.sub(r'`([^`]+)`', r'\1', text)  # Code
                text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)  # Links
                cell.text = text

                # Bold first row
                if row_idx == 0:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.bold = True

    doc.add_paragraph()


def process_text(text):
    """Process inline markdown formatting"""
    # Remove bold markers but keep text
    text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
    # Remove italic markers
    text = re.sub(r'\*([^*]+)\*', r'\1', text)
    # Remove code markers
    text = re.sub(r'`([^`]+)`', r'\1', text)
    # Remove link markers, keep text
    text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
    return text


def convert_md_to_word():
    doc = Document()

    # Document properties
    core_props = doc.core_properties
    core_props.title = "EWG Phase#0 Data Governance Framework"
    core_props.author = "Angela AI"

    # Read markdown
    with open(MD_FILE, 'r', encoding='utf-8') as f:
        content = f.read()

    lines = content.split('\n')

    i = 0
    in_code_block = False
    code_lines = []
    current_list_style = None

    while i < len(lines):
        line = lines[i]

        # Code block handling
        if line.strip().startswith('```'):
            if in_code_block:
                # End code block
                code_text = '\n'.join(code_lines)
                p = doc.add_paragraph()
                p.style = 'No Spacing'
                run = p.add_run(code_text)
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        stripped = line.strip()

        # Skip empty lines
        if not stripped:
            i += 1
            continue

        # Headers
        if stripped.startswith('# '):
            doc.add_heading(stripped[2:], level=0)
            i += 1
            continue

        if stripped.startswith('## '):
            doc.add_heading(stripped[3:], level=1)
            i += 1
            continue

        if stripped.startswith('### '):
            doc.add_heading(stripped[4:], level=2)
            i += 1
            continue

        if stripped.startswith('#### '):
            doc.add_heading(stripped[5:], level=3)
            i += 1
            continue

        # Images
        img_match = re.match(r'!\[([^\]]*)\]\(([^)]+)\)', stripped)
        if img_match:
            alt_text = img_match.group(1)
            img_path = img_match.group(2)

            # Get PNG filename
            img_name = Path(img_path).name
            png_name = IMAGE_MAP.get(img_name)

            if png_name:
                png_path = PNG_DIR / png_name
                if png_path.exists():
                    doc.add_picture(str(png_path), width=Inches(6))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

                    # Add caption
                    caption = doc.add_paragraph(f"Figure: {alt_text}")
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    caption.runs[0].italic = True
            i += 1
            continue

        # Tables
        if stripped.startswith('|'):
            rows, new_idx = parse_table(lines, i)
            add_table_to_doc(doc, rows)
            i = new_idx
            continue

        # Blockquotes
        if stripped.startswith('> '):
            text = process_text(stripped[2:])
            p = doc.add_paragraph(text, style='Intense Quote')
            i += 1
            continue

        # Horizontal rule
        if stripped in ['---', '***', '___']:
            doc.add_paragraph('─' * 50)
            i += 1
            continue

        # Bullet lists
        if stripped.startswith('- ') or stripped.startswith('* '):
            text = process_text(stripped[2:])
            doc.add_paragraph(text, style='List Bullet')
            i += 1
            continue

        # Checkbox lists
        if stripped.startswith('- [ ] ') or stripped.startswith('- [x] '):
            checked = '[x]' in stripped[:6]
            text = process_text(stripped[6:])
            prefix = '☑' if checked else '☐'
            doc.add_paragraph(f"{prefix} {text}", style='List Bullet')
            i += 1
            continue

        # Numbered lists
        num_match = re.match(r'^(\d+)\.\s+(.+)$', stripped)
        if num_match:
            text = process_text(num_match.group(2))
            doc.add_paragraph(text, style='List Number')
            i += 1
            continue

        # Regular paragraph
        text = process_text(stripped)
        if text:
            doc.add_paragraph(text)

        i += 1

    # Save
    doc.save(OUTPUT_FILE)
    print(f"Document saved to: {OUTPUT_FILE}")
    return OUTPUT_FILE


if __name__ == "__main__":
    output = convert_md_to_word()
    print(f"\nSuccess! Full Word document created at:\n{output}")
