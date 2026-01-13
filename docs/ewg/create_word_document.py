#!/usr/bin/env python3
"""
Create MS Word document summarizing EWG Phase#0 Data Governance Framework
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from pathlib import Path

# Paths
BASE_DIR = Path(__file__).parent
PNG_DIR = BASE_DIR / "diagrams" / "png"
OUTPUT_FILE = BASE_DIR / "EWG_Phase0_Data_Governance_Summary.docx"


def create_document():
    doc = Document()

    # Set document properties
    core_props = doc.core_properties
    core_props.title = "EWG Phase#0 Data Governance Framework Summary"
    core_props.author = "Angela AI"

    # ============================================
    # TITLE PAGE
    # ============================================
    title = doc.add_heading("EWG Phase#0", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_heading("Data Governance Framework Summary", level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_paragraph("Document Version: 1.0").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Date: January 13, 2026").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Prepared by: Angela AI").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("For: C-Level Executives, Data Management Team").alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # ============================================
    # TABLE OF CONTENTS
    # ============================================
    doc.add_heading("Table of Contents", level=1)

    toc_items = [
        "1. Executive Summary",
        "2. Data Governance Framework",
        "3. Data Policy Framework",
        "4. Time of Data Capture",
        "5. Period-End Close Policy (Listed Company)",
        "6. Water Volume Data Consolidation (SSOT)",
        "7. DAMA-DMBOK Reference Framework",
        "8. Implementation Roadmap",
    ]

    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.left_indent = Inches(0.5)

    doc.add_page_break()

    # ============================================
    # SECTION 1: EXECUTIVE SUMMARY
    # ============================================
    doc.add_heading("1. Executive Summary", level=1)

    doc.add_heading("1.1 Purpose", level=2)
    doc.add_paragraph(
        "This document establishes the framework for implementing Data Governance at EWG, "
        "addressing three critical areas identified during C-Level interviews:"
    )

    bullets = [
        ("Data Governance", "Authority and decision-making over data assets"),
        ("Data Policy", "Rules and standards for data handling"),
        ("Time of Data Capture", "Data freshness and timeliness requirements"),
    ]

    for title, desc in bullets:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(f"{title}: ")
        run.bold = True
        p.add_run(desc)

    doc.add_heading("1.2 Key Outcomes for Phase#0", level=2)
    outcomes = [
        "Establish governance structure with clear roles and responsibilities",
        "Define policy framework covering quality, security, access, and retention",
        "Document timeliness requirements for critical data domains",
        "Create roadmap for governance implementation",
        "Address water volume data consolidation issue (SSOT)",
    ]
    for outcome in outcomes:
        doc.add_paragraph(outcome, style='List Bullet')

    doc.add_page_break()

    # ============================================
    # SECTION 2: DATA GOVERNANCE FRAMEWORK
    # ============================================
    doc.add_heading("2. Data Governance Framework", level=1)

    doc.add_heading("2.1 Definition", level=2)
    p = doc.add_paragraph()
    p.add_run("Data Governance").bold = True
    p.add_run(" is the exercise of authority, control, and shared decision-making "
              "over the management of data assets. It ensures data is treated as a "
              "strategic enterprise asset with proper accountability.")

    doc.add_heading("2.2 Core Objectives", level=2)
    objectives = [
        ("Data Quality & Reliability", "Ensure accuracy, completeness, consistency"),
        ("Security & Privacy", "Protect sensitive information, risk mitigation"),
        ("Regulatory Compliance", "Maintain auditable controls, avoid penalties"),
        ("Decision Enablement", "Improve data discoverability for faster decisions"),
        ("Operational Efficiency", "Reduce redundant efforts, cost savings"),
    ]
    for i, (obj, desc) in enumerate(objectives, 1):
        p = doc.add_paragraph()
        p.add_run(f"{i}. {obj}: ").bold = True
        p.add_run(desc)

    doc.add_heading("2.3 The Four Pillars", level=2)
    doc.add_paragraph("Data Governance rests on four foundational pillars:")

    pillars = [
        ("People", "Governance Council, Data Owners, Data Stewards, Data Custodians"),
        ("Process", "Lifecycle Management, Issue Resolution, Change Management, QA"),
        ("Technology", "Data Catalog, Lineage Mapping, Access Controls, Quality Monitoring"),
        ("Policy", "Classification, Quality Standards, Access Rules, Retention Rules"),
    ]
    for pillar, details in pillars:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f"{pillar}: ").bold = True
        p.add_run(details)

    # Add diagram
    doc.add_paragraph()
    if (PNG_DIR / "01_data_governance_4_pillars.png").exists():
        doc.add_picture(str(PNG_DIR / "01_data_governance_4_pillars.png"), width=Inches(6))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("Figure 2.1: Data Governance Four Pillars").alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("2.4 Key Roles (RACI)", level=2)

    # Create RACI table
    table = doc.add_table(rows=5, cols=5)
    table.style = 'Table Grid'

    headers = ["Role", "Strategy", "Policies", "Quality Rules", "Implementation"]
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    raci_data = [
        ["Governance Council", "A", "A", "C", "I"],
        ["Data Owner", "C", "R", "A", "C"],
        ["Data Steward", "I", "C", "R", "C"],
        ["Data Custodian", "I", "I", "I", "R/A"],
    ]
    for row_idx, row_data in enumerate(raci_data, 1):
        for col_idx, cell_data in enumerate(row_data):
            table.rows[row_idx].cells[col_idx].text = cell_data

    doc.add_paragraph()
    doc.add_paragraph("R = Responsible, A = Accountable, C = Consulted, I = Informed",
                     style='Intense Quote')

    doc.add_page_break()

    # ============================================
    # SECTION 3: DATA POLICY FRAMEWORK
    # ============================================
    doc.add_heading("3. Data Policy Framework", level=1)

    doc.add_heading("3.1 Policy Hierarchy", level=2)
    hierarchy = [
        ("Level 1 - Enterprise Data Policy", "High-level principles and strategic direction"),
        ("Level 2 - Data Standards", "Technical specifications for implementation"),
        ("Level 3 - Data Procedures", "Step-by-step operational instructions"),
        ("Level 4 - Data Guidelines", "Best practices and recommendations"),
    ]
    for level, desc in hierarchy:
        p = doc.add_paragraph(style='List Number')
        p.add_run(f"{level}: ").bold = True
        p.add_run(desc)

    # Add diagram
    doc.add_paragraph()
    if (PNG_DIR / "02_data_policy_hierarchy.png").exists():
        doc.add_picture(str(PNG_DIR / "02_data_policy_hierarchy.png"), width=Inches(6))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("Figure 3.1: Data Policy Hierarchy").alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("3.2 Core Policy Domains", level=2)

    policies = [
        ("Data Classification Policy", [
            "Public - No restrictions on disclosure",
            "Internal - For business use only",
            "Confidential - Sensitive business information (need-to-know, encryption)",
            "Restricted - Highly sensitive, regulated (strict controls, audit logging)",
        ]),
        ("Data Quality Policy", [
            "Accuracy: Data correctly represents reality (< 1% error rate)",
            "Completeness: Required data is present (> 99%)",
            "Consistency: Data agrees across systems (0 discrepancy)",
            "Timeliness: Data is current when needed (Per SLA)",
            "Validity: Data conforms to rules (100%)",
            "Uniqueness: No unwanted duplicates (< 0.1%)",
        ]),
        ("Data Access Policy", [
            "Open: All employees can access public data",
            "Controlled: Department members with manager approval",
            "Restricted: Specific roles with data owner approval",
            "Privileged: Named individuals with governance council approval",
        ]),
        ("Data Retention Policy", [
            "Transactional: 2 years active + 5 years archive = 7 years total",
            "Customer PII: Active relationship + 3 years (per consent)",
            "Financial Records: 1 year active + 9 years archive = 10 years total",
            "Audit Logs: 1 year active + 6 years archive = 7 years total",
        ]),
        ("Data Security Policy", [
            "Encryption at Rest: AES-256 for Confidential+",
            "Encryption in Transit: TLS 1.2+ required",
            "Access Authentication: MFA for Confidential+",
            "Audit Logging: All access logged",
        ]),
    ]

    for policy_name, items in policies:
        doc.add_heading(policy_name, level=3)
        for item in items:
            doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    # ============================================
    # SECTION 4: TIME OF DATA CAPTURE
    # ============================================
    doc.add_heading("4. Time of Data Capture", level=1)

    doc.add_heading("4.1 Key Concepts", level=2)
    concepts = [
        ("Data Timeliness", "Degree to which data is available when needed"),
        ("Data Freshness", "Age of data at any given point (Current Time - Last Update)"),
        ("Data Latency", "Delay between generation and availability"),
        ("Data Currency", "How current relative to real-world state"),
        ("Data Volatility", "Rate at which data changes"),
    ]
    for concept, desc in concepts:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f"{concept}: ").bold = True
        p.add_run(desc)

    # Add diagram
    doc.add_paragraph()
    if (PNG_DIR / "03_time_of_data_capture_flow.png").exists():
        doc.add_picture(str(PNG_DIR / "03_time_of_data_capture_flow.png"), width=Inches(6))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("Figure 4.1: Time of Data Capture Flow").alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("4.2 Critical Timestamps", level=2)
    timestamps = [
        ("event_timestamp", "When the event actually occurred"),
        ("created_at", "When record was created in source"),
        ("extracted_at", "When data was extracted (ETL)"),
        ("loaded_at", "When data was loaded to destination"),
        ("updated_at", "When record was last modified"),
        ("valid_from / valid_to", "Temporal validity window"),
    ]
    for ts, desc in timestamps:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f"{ts}: ").bold = True
        p.add_run(desc)

    doc.add_heading("4.3 Timeliness Requirements", level=2)

    # Create timeliness table
    table = doc.add_table(rows=6, cols=4)
    table.style = 'Table Grid'

    headers = ["Data Domain", "Freshness", "Latency", "Update Frequency"]
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    timeliness_data = [
        ["Real-time Analytics", "< 1 min", "Seconds", "Continuous"],
        ["Operational Reporting", "< 1 hour", "Minutes", "Near real-time"],
        ["Daily Reporting", "< 24 hours", "Hours", "Daily batch"],
        ["Historical Analysis", "< 1 week", "Days", "Weekly batch"],
        ["Regulatory Reporting", "Per regulation", "As specified", "Per deadline"],
    ]
    for row_idx, row_data in enumerate(timeliness_data, 1):
        for col_idx, cell_data in enumerate(row_data):
            table.rows[row_idx].cells[col_idx].text = cell_data

    doc.add_page_break()

    # ============================================
    # SECTION 5: PERIOD-END CLOSE POLICY
    # ============================================
    doc.add_heading("5. Period-End Close Policy (Listed Company)", level=1)

    p = doc.add_paragraph()
    p.add_run("Critical for EWG: ").bold = True
    p.add_run("As a SET-listed company, EWG must comply with strict regulatory "
              "reporting requirements from SET and SEC Thailand.")

    doc.add_heading("5.1 Regulatory Requirements", level=2)
    requirements = [
        ("Quarterly Financial Statements (Reviewed)", "45 days from quarter-end"),
        ("Annual Financial Statements (Audited)", "90 days from fiscal year-end"),
        ("Form 56-1 One Report", "3 months from fiscal year-end"),
    ]
    for req, deadline in requirements:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f"{req}: ").bold = True
        p.add_run(deadline)

    doc.add_heading("5.2 Penalties for Non-Compliance", level=2)
    penalties = [
        "Late submission: THB 100,000 + THB 3,000/day",
        "Incomplete information: THB 100,000 + THB 3,000/day",
        "Director share reporting: Up to THB 500,000 + THB 10,000/day",
    ]
    for penalty in penalties:
        doc.add_paragraph(penalty, style='List Bullet')

    doc.add_heading("5.3 Close Calendar Overview", level=2)
    close_types = [
        ("Month-End (Soft Close)", "Reconciliations, accruals, basic reporting"),
        ("Quarter-End (Hard Close)", "Full reconciliation, external reporting, MD&A"),
        ("Year-End (Full Close)", "Audit preparation, tax adjustments, Form 56-1"),
    ]
    for close_type, activities in close_types:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f"{close_type}: ").bold = True
        p.add_run(activities)

    # Add diagram
    doc.add_paragraph()
    if (PNG_DIR / "05_period_end_close_timeline.png").exists():
        doc.add_picture(str(PNG_DIR / "05_period_end_close_timeline.png"), width=Inches(6))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("Figure 5.1: Period-End Close Timeline").alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("5.4 Data Freeze Policy", level=2)
    freeze_levels = [
        ("Level 1: Soft Freeze", "No new transactions in source systems (WD 0-2)"),
        ("Level 2: Hard Freeze", "No changes to GL without approval (WD 3-10)"),
        ("Level 3: Audit Freeze", "Only audit adjustments allowed (WD 11-Filing)"),
        ("Level 4: Archive", "Data locked, read-only (Post-filing)"),
    ]
    for level, desc in freeze_levels:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f"{level}: ").bold = True
        p.add_run(desc)

    # Add diagram
    doc.add_paragraph()
    if (PNG_DIR / "06_data_freeze_policy.png").exists():
        doc.add_picture(str(PNG_DIR / "06_data_freeze_policy.png"), width=Inches(6))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("Figure 5.2: Data Freeze Policy & Exception Workflow").alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("5.5 Data Quality Gates", level=2)
    gates = [
        ("G1 - Completeness (WD+1)", "All sources loaded"),
        ("G2 - Accuracy (WD+2)", "Reconciliation pass"),
        ("G3 - Timeliness (WD+3)", "Data within SLA"),
        ("G4 - Consistency (WD+4)", "Cross-system match"),
        ("G5 - Sign-off (WD+5)", "Steward approval"),
    ]
    for gate, check in gates:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f"{gate}: ").bold = True
        p.add_run(check)

    doc.add_page_break()

    # ============================================
    # SECTION 6: WATER VOLUME DATA CONSOLIDATION
    # ============================================
    doc.add_heading("6. Water Volume Data Consolidation (SSOT)", level=1)

    p = doc.add_paragraph()
    p.add_run("Issue Identified: ").bold = True
    p.add_run("Water volume data is not consolidated across the organization. "
              "Different departments use different data sources, leading to "
              "inconsistent reporting and decision-making challenges.")

    doc.add_heading("6.1 Pain Points", level=2)
    pain_points = [
        "Multiple Data Sources - No single version of truth",
        "Manual Data Entry - High error rate, delays",
        "No Standard Definitions - Apples-to-oranges comparisons",
        "Delayed Reporting - Decisions based on stale data",
        "Reconciliation Overhead - 20-30% of analyst time wasted",
        "Audit Challenges - Cannot trace data lineage",
    ]
    for point in pain_points:
        doc.add_paragraph(point, style='List Bullet')

    doc.add_heading("6.2 Target State: Single Source of Truth (SSOT)", level=2)
    p = doc.add_paragraph()
    p.add_run("\"One number, one source, one truth - regardless of who asks or when they ask.\"")
    p.italic = True

    # Add diagram
    doc.add_paragraph()
    if (PNG_DIR / "07_ssot_architecture.png").exists():
        doc.add_picture(str(PNG_DIR / "07_ssot_architecture.png"), width=Inches(6))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("Figure 6.1: SSOT Architecture").alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("6.3 Golden Record Framework", level=2)
    doc.add_paragraph("A Golden Record is the single, authoritative version of a data entity.")

    doc.add_heading("Water Volume Data Elements:", level=3)
    elements = [
        ("Raw Water Intake", "Volume entering treatment (m³) - Source: SCADA"),
        ("Treated Water Production", "Volume after treatment (m³) - Source: SCADA"),
        ("Water Distribution", "Volume to distribution network (m³) - Source: SCADA"),
        ("Water Sales", "Volume billed to customers (m³) - Source: Billing"),
        ("Non-Revenue Water (NRW)", "Water lost (leakage, theft) (m³) - Calculated"),
    ]
    for element, desc in elements:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f"{element}: ").bold = True
        p.add_run(desc)

    doc.add_heading("6.4 Data Consolidation Roadmap", level=2)

    # Add diagram
    if (PNG_DIR / "08_data_consolidation_roadmap.png").exists():
        doc.add_picture(str(PNG_DIR / "08_data_consolidation_roadmap.png"), width=Inches(6))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("Figure 6.2: Data Consolidation Roadmap").alignment = WD_ALIGN_PARAGRAPH.CENTER

    phases = [
        ("Phase 1: Foundation (Month 1-2)", [
            "Data Source Inventory",
            "Data Element Mapping",
            "Quality Assessment",
            "Stakeholder Alignment",
        ]),
        ("Phase 2: Integration (Month 3-4)", [
            "ETL/ELT Pipeline Development",
            "Validation Rules Implementation",
            "Master Data Repository Setup",
            "Initial Data Load & Reconciliation",
        ]),
        ("Phase 3: Adoption (Month 5-6)", [
            "Dashboard Development",
            "Report Migration",
            "Training & Change Management",
            "Go-Live & Support",
        ]),
        ("Phase 4: Optimization (Ongoing)", [
            "Continuous Quality Monitoring",
            "Feedback Loop",
            "Process Improvement",
            "Audit Support",
        ]),
    ]

    for phase_name, activities in phases:
        doc.add_heading(phase_name, level=3)
        for activity in activities:
            doc.add_paragraph(activity, style='List Bullet')

    doc.add_heading("6.5 Success Metrics", level=2)
    metrics = [
        ("Report discrepancy rate", "Baseline: High → Target: 0%"),
        ("Time to produce reports", "Baseline: 3-5 days → Target: < 1 day"),
        ("Analyst time on reconciliation", "Baseline: 20-30% → Target: < 5%"),
        ("Data freshness", "Baseline: Days old → Target: < 1 hour"),
    ]
    for metric, target in metrics:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f"{metric}: ").bold = True
        p.add_run(target)

    doc.add_page_break()

    # ============================================
    # SECTION 7: DAMA-DMBOK
    # ============================================
    doc.add_heading("7. DAMA-DMBOK Reference Framework", level=1)

    doc.add_paragraph(
        "The DAMA Data Management Body of Knowledge (DMBOK) is the globally recognized "
        "standard for data management. It positions Data Governance at the center of "
        "11 interconnected knowledge areas."
    )

    # Add diagram
    if (PNG_DIR / "04_dama_wheel.png").exists():
        doc.add_picture(str(PNG_DIR / "04_dama_wheel.png"), width=Inches(5.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("Figure 7.1: DAMA Wheel - 11 Knowledge Areas").alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("7.1 The 11 Knowledge Areas", level=2)
    knowledge_areas = [
        ("Data Governance", "Central orchestrating function (Policies, RACI)"),
        ("Data Architecture", "Blueprint for data structures"),
        ("Data Modeling & Design", "Conceptual, logical, physical models"),
        ("Data Storage & Operations", "Database management, backup"),
        ("Data Security", "Access control, encryption"),
        ("Data Integration & Interoperability", "ETL/ELT, APIs"),
        ("Document & Content Management", "Unstructured data handling"),
        ("Data Warehousing & BI", "Analytical data management"),
        ("Metadata Management", "Data catalog, lineage"),
        ("Reference & Master Data Management", "MDM, Golden Records"),
        ("Data Quality Management", "Profiling, cleansing"),
    ]
    for i, (area, desc) in enumerate(knowledge_areas, 1):
        p = doc.add_paragraph()
        p.add_run(f"{i}. {area}: ").bold = True
        p.add_run(desc)

    doc.add_page_break()

    # ============================================
    # SECTION 8: IMPLEMENTATION ROADMAP
    # ============================================
    doc.add_heading("8. Implementation Roadmap", level=1)

    doc.add_heading("8.1 Phase#0 Work Packages", level=2)

    work_packages = [
        ("WP1", "Current State Assessment", "Week 1-2", "Consultant"),
        ("WP2", "Governance Structure Design", "Week 2-3", "CDO"),
        ("WP3", "Policy Framework Development", "Week 3-4", "Data Governance Lead"),
        ("WP4", "Data Inventory", "Week 3-5", "Data Stewards"),
        ("WP5", "Timeliness Requirements", "Week 4-5", "Data Owners"),
        ("WP6", "Quick Wins Identification", "Week 5-6", "Project Team"),
    ]

    # Create work packages table
    table = doc.add_table(rows=len(work_packages)+1, cols=4)
    table.style = 'Table Grid'

    headers = ["WP#", "Work Package", "Duration", "Owner"]
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    for row_idx, row_data in enumerate(work_packages, 1):
        for col_idx, cell_data in enumerate(row_data):
            table.rows[row_idx].cells[col_idx].text = cell_data

    doc.add_heading("8.2 Quick Wins for Phase#0", level=2)
    quick_wins = [
        "Define \"Official\" Water Volume (Low effort, High impact)",
        "Identify authoritative source system (Low effort, High impact)",
        "Create simple reconciliation report (Medium effort, High impact)",
        "Establish monthly data review meeting (Low effort, Medium impact)",
        "Document current data flow (Medium effort, Medium impact)",
    ]
    for i, win in enumerate(quick_wins, 1):
        p = doc.add_paragraph()
        p.add_run(f"{i}. ").bold = True
        p.add_run(win)

    doc.add_heading("8.3 Key Questions for C-Level", level=2)

    questions = [
        ("Data Governance", [
            "Who should own which data domains?",
            "What governance structure fits EWG's culture?",
            "How will governance decisions be escalated?",
        ]),
        ("Data Policy", [
            "What regulatory requirements apply (PDPA)?",
            "How should data be classified?",
            "What are the retention requirements?",
        ]),
        ("Time of Data Capture", [
            "What are the freshness requirements per function?",
            "Where are the critical latency bottlenecks?",
            "What is the acceptable data delay for decisions?",
        ]),
    ]

    for category, items in questions:
        doc.add_heading(category, level=3)
        for item in items:
            doc.add_paragraph(item, style='List Bullet')

    doc.add_heading("8.4 Success Metrics", level=2)
    success_metrics = [
        ("Stakeholder interviews completed", "100% of C-Level"),
        ("Critical data elements identified", "Top 20 CDEs"),
        ("Policy gaps documented", "Complete assessment"),
        ("Quick wins identified", "Minimum 5"),
        ("Governance structure approved", "Signed charter"),
    ]
    for metric, target in success_metrics:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f"{metric}: ").bold = True
        p.add_run(target)

    # ============================================
    # FOOTER
    # ============================================
    doc.add_page_break()
    doc.add_heading("Document Information", level=1)

    info = [
        ("Prepared by:", "Angela AI"),
        ("For:", "EWG Phase#0 Data Governance Initiative"),
        ("Date:", "January 13, 2026"),
        ("Version:", "1.0"),
    ]
    for label, value in info:
        p = doc.add_paragraph()
        p.add_run(f"{label} ").bold = True
        p.add_run(value)

    # Save document
    doc.save(OUTPUT_FILE)
    print(f"Document saved to: {OUTPUT_FILE}")
    return OUTPUT_FILE


if __name__ == "__main__":
    output = create_document()
    print(f"\nSuccess! Word document created at:\n{output}")
